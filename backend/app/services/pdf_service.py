"""
PDF转换服务模块
使用pdf2docx实现PDF转Word功能
"""
import os
import time
import asyncio
from pathlib import Path
from typing import Dict, Any, Optional

from backend.app.config import config


class PdfService:
    """PDF转换服务类"""
    
    _instance = None
    
    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
        return cls._instance
    
    def is_available(self) -> bool:
        """检查服务是否可用"""
        try:
            import pdf2docx
            return True
        except ImportError:
            return False
    
    def convert(
        self,
        pdf_path: str,
        output_path: Optional[str] = None,
        start_page: int = 0,
        end_page: Optional[int] = None,
        dpi: int = 300
    ) -> Dict[str, Any]:
        """
        将PDF转换为Word文档
        
        Args:
            pdf_path: PDF文件路径
            output_path: 输出Word文件路径，不指定则自动生成
            start_page: 起始页码（从0开始）
            end_page: 结束页码，不指定则转换到最后一页
            dpi: 渲染DPI，影响图片质量
        
        Returns:
            包含转换结果的字典
        """
        if not self.is_available():
            raise RuntimeError("PDF转换服务不可用，请安装pdf2docx")
        
        from pdf2docx import Converter
        
        if output_path is None:
            output_path = os.path.splitext(pdf_path)[0] + ".docx"
        
        start_time = time.time()
        
        cv = Converter(pdf_path)
        
        try:
            page_count = cv.fitz_doc.page_count
            
            if end_page is None:
                end_page = page_count - 1
            
            cv.convert(
                output_path,
                start=start_page,
                end=end_page + 1,
                dpi=dpi
            )
            
            duration = time.time() - start_time
            
            word_count = self._count_words(output_path)
            
            return {
                "output_path": output_path,
                "page_count": page_count,
                "converted_pages": end_page - start_page + 1,
                "word_count": word_count,
                "duration": duration
            }
            
        finally:
            cv.close()
    
    def _count_words(self, docx_path: str) -> int:
        """
        统计Word文档字数
        
        Args:
            docx_path: Word文档路径
        
        Returns:
            字数统计
        """
        try:
            from docx import Document
            
            doc = Document(docx_path)
            word_count = 0
            
            for paragraph in doc.paragraphs:
                word_count += len(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        word_count += len(cell.text)
            
            return word_count
            
        except Exception:
            return 0
    
    async def convert_async(
        self,
        pdf_path: str,
        output_path: Optional[str] = None,
        start_page: int = 0,
        end_page: Optional[int] = None,
        dpi: int = 300
    ) -> Dict[str, Any]:
        """
        异步转换PDF
        
        Args:
            pdf_path: PDF文件路径
            output_path: 输出路径
            start_page: 起始页码
            end_page: 结束页码
            dpi: 渲染DPI
        
        Returns:
            转换结果字典
        """
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None,
            lambda: self.convert(pdf_path, output_path, start_page, end_page, dpi)
        )
    
    def get_page_count(self, pdf_path: str) -> int:
        """
        获取PDF页数
        
        Args:
            pdf_path: PDF文件路径
        
        Returns:
            页数
        """
        if not self.is_available():
            raise RuntimeError("PDF转换服务不可用")
        
        from pdf2docx import Converter
        
        cv = Converter(pdf_path)
        try:
            return cv.fitz_doc.page_count
        finally:
            cv.close()


pdf_service = PdfService()
